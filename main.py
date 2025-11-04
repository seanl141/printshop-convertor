from fastapi import FastAPI, File, UploadFile, Form
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi import Request
from fastapi import Form 
from fastapi.responses import JSONResponse
import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
from docx2pdf import convert
from pdf2docx import Converter
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
from pillow_heif import register_heif_opener
from PIL import Image
import shutil

register_heif_opener()

app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

converted_dir = Path("converted")
converted_dir.mkdir(exist_ok=True)

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/convert_docx")
async def convert_docx(file: UploadFile = File(...)):
    input_path = Path("temp_" + file.filename)
    with open(input_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
    convert(input_path)
    output_path = input_path.with_suffix(".pdf")
    dest = converted_dir / output_path.name
    shutil.move(output_path, dest)
    input_path.unlink(missing_ok=True)
    return FileResponse(dest, filename=dest.name)

@app.post("/convert_heic")
async def convert_heic(file: UploadFile = File(...)):
    input_path = Path("temp_" + file.filename)
    with open(input_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
    img = Image.open(input_path)
    output_path = converted_dir / (input_path.stem + ".jpg")
    img.save(output_path, "JPEG")
    input_path.unlink(missing_ok=True)
    return FileResponse(output_path, filename=output_path.name)

@app.post("/merge_pdfs")
async def merge_pdfs(files: list[UploadFile] = File(...)):
    merger = PdfMerger()
    paths = []
    for f in files:
        p = Path("temp_" + f.filename)
        with open(p, "wb") as out:
            shutil.copyfileobj(f.file, out)
        merger.append(str(p))
        paths.append(p)
    output = converted_dir / "merged.pdf"
    merger.write(output)
    merger.close()
    for p in paths:
        p.unlink(missing_ok=True)
    return FileResponse(output, filename="merged.pdf")

@app.post("/split_pdf")
async def split_pdf(file: UploadFile = File(...), pages: str = Form("")):
    input_path = Path("temp_" + file.filename)
    with open(input_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    reader = PdfReader(input_path)
    total_pages = len(reader.pages)

    # Parse the page input (e.g., "1-3,5,8")
    selected_pages = set()
    if pages.strip():
        parts = pages.split(",")
        for part in parts:
            if "-" in part:
                start, end = part.split("-")
                selected_pages.update(range(int(start), int(end) + 1))
            else:
                selected_pages.add(int(part))
    else:
        selected_pages = set(range(1, total_pages + 1))

    writer = PdfWriter()
    for page_num in sorted(selected_pages):
        if 1 <= page_num <= total_pages:
            writer.add_page(reader.pages[page_num - 1])

    output_path = converted_dir / f"{input_path.stem}_selected.pdf"
    with open(output_path, "wb") as out:
        writer.write(out)

    input_path.unlink(missing_ok=True)
    return FileResponse(output_path, filename=output_path.name)

@app.post("/pdf_to_word")
async def pdf_to_word(file: UploadFile = File(...)):
    try:
        # Save uploaded PDF temporarily
        input_path = Path("temp_" + file.filename)
        with open(input_path, "wb") as f:
            shutil.copyfileobj(file.file, f)

        output_path = converted_dir / (input_path.stem + ".docx")

        # Open PDF
        pdf = fitz.open(input_path)
        docx = Document()

        # Loop through each page, extract text, add to Word
        for i, page in enumerate(pdf, start=1):
            text = page.get_text("text")
            if not text.strip():
                text = "[No extractable text on this page]"
            docx.add_heading(f"Page {i}", level=2)
            docx.add_paragraph(text)
            if i < len(pdf):
                docx.add_page_break()

        # Save the Word file
        docx.save(output_path)
        pdf.close()
        input_path.unlink(missing_ok=True)

        # Return the converted .docx file
        return FileResponse(output_path, filename=output_path.name)

    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"PDF → Word conversion failed: {str(e)}"}
        )
    
    #leaflets and business cards

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from PyPDF2 import PdfReader
from PIL import Image
import fitz
import shutil
from fastapi.responses import FileResponse, JSONResponse
from pathlib import Path

@app.post("/layout_cards")
async def layout_cards(
    front: UploadFile = File(...),
    back: UploadFile = File(...),
    product: str = Form("business_card")
):
    try:
        # --- Spec definitions ---
        SIZES = {
            "business_card": {"w": 85, "h": 55, "rows": 4, "cols": 2},
            "a5_leaflet": {"w": 148, "h": 210, "rows": 1, "cols": 2},
        }
        spec = SIZES.get(product, SIZES["business_card"])
        w, h, rows, cols = spec["w"], spec["h"], spec["rows"], spec["cols"]

        # --- Save uploaded files temporarily ---
        front_path = Path("temp_front_" + front.filename)
        back_path = Path("temp_back_" + back.filename)
        with open(front_path, "wb") as f:
            shutil.copyfileobj(front.file, f)
        with open(back_path, "wb") as f:
            shutil.copyfileobj(back.file, f)

        output_path = converted_dir / f"layout_{product}.pdf"
        c = canvas.Canvas(str(output_path), pagesize=A4)
        page_w, page_h = A4

        # --- Helper function to draw a side (image or PDF) ---
        def draw_side(file_path: Path):
            ext = file_path.suffix.lower()

            if ext in [".jpg", ".jpeg", ".png"]:
                target_w, target_h = w * mm, h * mm
                for r in range(rows):
                    for col in range(cols):
                        x = col * target_w
                        y = page_h - (r + 1) * target_h
                        c.drawImage(str(file_path), x, y, width=target_w, height=target_h)
                        # crop marks
                        c.setLineWidth(0.2)
                        c.line(x, y, x + 10, y)
                        c.line(x, y, x, y + 10)
                        c.line(x + target_w, y, x + target_w - 10, y)
                        c.line(x + target_w, y, x + target_w, y + 10)
                        c.line(x, y + target_h, x + 10, y + target_h)
                        c.line(x, y + target_h, x, y + target_h - 10)
                        c.line(x + target_w, y + target_h, x + target_w - 10, y + target_h)
                        c.line(x + target_w, y + target_h, x + target_w, y + target_h - 10)

            elif ext == ".pdf":
                # Render first page of the PDF as image
                doc = fitz.open(file_path)
                if len(doc) == 0:
                    raise Exception("Empty PDF")
                pix = doc[0].get_pixmap(dpi=300)
                img_path_temp = file_path.with_suffix(".jpg")
                pix.save(img_path_temp)
                doc.close()
                draw_side(img_path_temp)
                img_path_temp.unlink(missing_ok=True)

            else:
                raise Exception("Unsupported file type (use .pdf, .jpg, or .png)")

        # --- Front page ---
        draw_side(front_path)
        c.showPage()

        # --- Back page ---
        draw_side(back_path)
        c.showPage()

        c.save()
        front_path.unlink(missing_ok=True)
        back_path.unlink(missing_ok=True)

        return FileResponse(output_path, filename=output_path.name)

    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})